from docx import Document
from datetime import date
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def write(document_name, university_name, course_name, language, topics_names, topics_content, topics_timestamps,
          word_list):
    is_empty = False
    if not topics_names and not topics_content and not topics_timestamps:
        is_empty = True

    is_hebrew = False
    if language == "Hebrew":
        is_hebrew = True

    today = date.today()

    # Textual month, day and year
    today = today.strftime("%B %d, %Y")

    document = Document()

    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "{}\t{}\t{}".format(today, university_name, course_name)
    paragraph.style = document.styles["Header"]

    h = document.add_heading(document_name, 0)
    if is_hebrew:
        h.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    for topic, content, timestamp in zip(topics_names, topics_content, topics_timestamps):
        heading, topic_content = "", ""

        for word in topic:
            heading += (word + " ")

        for word in content:
            topic_content += (word + " ")

        if is_hebrew:
            heading += " נקודת הזמן בהקלטה: {} - {}".format(timestamp[0], timestamp[1])
        else:
            heading += " (Timestamp In Lecture: {} - {})".format(timestamp[0], timestamp[1])

        h = document.add_heading(heading, 1)
        p = document.add_paragraph(topic_content)

        for run in p.runs:
            run.font.name = 'Arial'

        if is_hebrew:
            h.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # can't divide the given topics according to timestamps, so just write the text in single paragraph:
    if is_empty:
        text = ""
        for word in word_list:
            text += word + " "

        p = document.add_paragraph(text)
        for run in p.runs:
            run.font.name = 'Arial'
        if is_hebrew:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    document.save('{}.docx'.format(document_name))
